import streamlit as st
import pandas as pd
from datetime import datetime
import io
import re
from openpyxl import load_workbook
from docx import Document

st.title("エクスポート機能")

# テンプレート管理タブ
tab1, tab2 = st.tabs(["テンプレート使用", "テンプレート登録"])

try:
    import sys
    import os
    sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))
    
    from utils.sheets_manager import SheetsManager
    from utils.template_manager import TemplateManager
    from utils.styles import MAIN_CSS, get_success_html, get_warning_html
    
    # CSS適用
    st.markdown(MAIN_CSS, unsafe_allow_html=True)
    
    @st.cache_resource
    def get_managers():
        return SheetsManager(), TemplateManager()
    
    sheets_manager, template_manager = get_managers()
    
    if not sheets_manager.is_connected():
        st.error("Google Sheets接続エラー")
        st.stop()
    
    # テンプレート変数の定義
    TEMPLATE_VARIABLES = {
        "債務者情報": {
            "{debtor_name}": "債務者名",
            "{court_name}": "裁判所名",
            "{case_number}": "事件番号",
            "{procedure_type}": "手続種別（個人再生/自己破産）",
            "{today}": "今日の日付（YYYY年MM月DD日）",
            "{today_slash}": "今日の日付（YYYY/MM/DD）",
            "{total_creditors}": "債権者総数",
            "{total_claim_amount}": "債権金額総計"
        },
        "債権者情報（n=1,2,3...で番号指定）": {
            "{id_n}": "ID",
            "{company_name_n}": "会社名",
            "{branch_name_n}": "支店名", 
            "{postal_code_n}": "郵便番号",
            "{address_n}": "住所",
            "{phone_number_n}": "電話番号",
            "{fax_number_n}": "FAX番号",
            "{claim_name_n}": "債権名",
            "{claim_amount_n}": "債権額",
            "{contract_date_n}": "契約日",
            "{first_borrowing_date_n}": "初回借入日",
            "{last_borrowing_date_n}": "最終借入日", 
            "{last_payment_date_n}": "最終返済日",
            "{original_creditor_n}": "原債権者",
            "{substitution_or_transfer_n}": "代位弁済/債権譲渡",
            "{transfer_date_n}": "債権移転日",
            "{status_n}": "ステータス",
            "{notes_n}": "備考",
            "{registration_date_n}": "登録日"
        },
        "計算・集計": {
            "{sum_claim_amount_1_to_n}": "1番目からn番目までの債権額合計",
            "{creditor_rank_n}": "n番目の債権者順位（1,2,3...）"
        }
    }
    
    def handle_dataframe_conversion(data):
        """DataFrameまたはリストをリスト形式に変換"""
        if data is None:
            return None, None
            
        # DataFrameの場合
        if isinstance(data, pd.DataFrame):
            if data.empty:
                return None, None
            headers = data.columns.tolist()
            creditor_data = data.to_dict('records')
            return headers, creditor_data
            
        # リストの場合
        elif isinstance(data, list) and len(data) > 1:
            headers = data[0]
            creditor_data = []
            for row in data[1:]:
                creditor_dict = {}
                for i, header in enumerate(headers):
                    creditor_dict[header] = row[i] if i < len(row) else ''
                creditor_data.append(creditor_dict)
            return headers, creditor_data
            
        return None, None

    def safe_get_spreadsheet_data_by_id(sheets_manager, spreadsheet_id):
        """スプレッドシートIDから安全にデータを取得"""
        try:
            data = sheets_manager.get_data_by_id(spreadsheet_id)
            
            if data is None or not data or len(data) <= 1:
                return None
            
            return data
            
        except Exception as e:
            st.error(f"データ取得中にエラーが発生しました: {e}")
            return None

    def safe_get_data_from_sheet_info(sheets_manager, sheet_info):
        """sheet_infoからデータを安全に取得"""
        try:
            # sheet_idキーがない場合は手動で追加
            if isinstance(sheet_info, dict):
                if 'sheet_id' not in sheet_info and 'id' in sheet_info:
                    sheet_info = sheet_info.copy()
                    sheet_info['sheet_id'] = sheet_info['id']
            
            data = sheets_manager.get_data(sheet_info)
            return data
            
        except Exception as e:
            st.error(f"データ取得中にエラーが発生しました: {e}")
            return None

    def replace_template_variables(text, creditor_data, debtor_name, court_name, procedure_type, case_number=""):
        """テンプレート変数を実際のデータで置換"""
        if not isinstance(text, str):
            return text
            
        result = text
        
        # 基本情報の置換
        result = result.replace("{debtor_name}", str(debtor_name))
        result = result.replace("{court_name}", str(court_name))
        result = result.replace("{case_number}", str(case_number))
        result = result.replace("{procedure_type}", str(procedure_type))
        result = result.replace("{today}", datetime.now().strftime('%Y年%m月%d日'))
        result = result.replace("{today_slash}", datetime.now().strftime('%Y/%m/%d'))
        result = result.replace("{total_creditors}", str(len(creditor_data)))
        
        # 債権金額総計の計算
        total_amount = sum(float(str(row.get('債権額', 0)).replace(',', '')) if row.get('債権額') else 0 for row in creditor_data)
        result = result.replace("{total_claim_amount}", f"{int(total_amount):,}")
        
        # 債権者個別情報の置換
        for i, creditor in enumerate(creditor_data, 1):
            replacements = {
                f"{{id_{i}}}": str(creditor.get('ID', '')),
                f"{{company_name_{i}}}": str(creditor.get('会社名', '')),
                f"{{branch_name_{i}}}": str(creditor.get('支店名', '')),
                f"{{postal_code_{i}}}": str(creditor.get('郵便番号', '')),
                f"{{address_{i}}}": str(creditor.get('住所', '')),
                f"{{phone_number_{i}}}": str(creditor.get('電話番号', '')),
                f"{{fax_number_{i}}}": str(creditor.get('FAX番号', '')),
                f"{{claim_name_{i}}}": str(creditor.get('債権名', '')),
                f"{{claim_amount_{i}}}": str(creditor.get('債権額', '')),
                f"{{contract_date_{i}}}": str(creditor.get('契約日', '')),
                f"{{first_borrowing_date_{i}}}": str(creditor.get('初回借入日', '')),
                f"{{last_borrowing_date_{i}}}": str(creditor.get('最終借入日', '')),
                f"{{last_payment_date_{i}}}": str(creditor.get('最終返済日', '')),
                f"{{original_creditor_{i}}}": str(creditor.get('原債権者', '')),
                f"{{substitution_or_transfer_{i}}}": str(creditor.get('代位弁済/債権譲渡', '')),
                f"{{transfer_date_{i}}}": str(creditor.get('債権移転日', '')),
                f"{{status_{i}}}": str(creditor.get('ステータス', '')),
                f"{{notes_{i}}}": str(creditor.get('備考', '')),
                f"{{registration_date_{i}}}": str(creditor.get('登録日', '')),
                f"{{creditor_rank_{i}}}": str(i)
            }
            
            for var, value in replacements.items():
                result = result.replace(var, value)
        
        return result
        
    def process_excel_template(template_path, creditor_data, debtor_name, court_name, procedure_type, case_number=""):
        """Excelテンプレートファイルを処理"""
        wb = load_workbook(template_path)
        
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value:
                        cell.value = replace_template_variables(
                            cell.value, creditor_data, debtor_name, court_name, procedure_type, case_number
                        )
        
        return wb
    
    def process_word_template(template_path, creditor_data, debtor_name, court_name, procedure_type, case_number=""):
        """Wordテンプレートファイルを処理"""
        doc = Document(template_path)
        
        # 段落内のテキストを処理
        for paragraph in doc.paragraphs:
            if paragraph.text:
                new_text = replace_template_variables(
                    paragraph.text, creditor_data, debtor_name, court_name, procedure_type, case_number
                )
                if new_text != paragraph.text:
                    paragraph.clear()
                    paragraph.add_run(new_text)
        
        # テーブル内のテキストを処理
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text:
                            new_text = replace_template_variables(
                                paragraph.text, creditor_data, debtor_name, court_name, procedure_type, case_number
                            )
                            if new_text != paragraph.text:
                                paragraph.clear()
                                paragraph.add_run(new_text)
        
        return doc
    
    def get_file_extension(file_path):
        """ファイル拡張子を取得"""
        return os.path.splitext(file_path)[1].lower()

    def get_template_key(court_name, procedure_type):
        """テンプレートキーを生成"""
        return f"{court_name}_{procedure_type}"

    # 裁判所と手続種別の選択肢
    courts = [
        "東京地方裁判所",
        "大阪地方裁判所", 
        "名古屋地方裁判所",
        "横浜地方裁判所",
        "神戸地方裁判所",
        "福岡地方裁判所",
        "仙台地方裁判所",
        "広島地方裁判所",
        "札幌地方裁判所",
        "その他"
    ]
    
    procedure_types = ["個人再生", "自己破産"]
    
    # テンプレート使用タブ
    with tab1:
        st.subheader("債権者一覧表テンプレート使用")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            selected_court = st.selectbox("裁判所を選択", courts)
            if selected_court == "その他":
                selected_court = st.text_input("裁判所名を入力")
        
        with col2:
            procedure_type = st.selectbox("手続種別を選択", procedure_types)
        
        with col3:
            case_number = st.text_input("事件番号（任意）", placeholder="例：令和6年(フ)第123号")
        
        template_key = f"{selected_court}_{procedure_type}"
        
        if template_manager.template_exists(template_key):
            template_info = template_manager.get_template_info(template_key)
            st.success(f"{selected_court} - {procedure_type} のテンプレートが利用可能です")
            st.text(f"説明: {template_info['description']}")
            st.text(f"最終更新: {template_info['last_modified']}")
            
            # 債務者データ取得・エクスポート機能
            st.markdown("---")
            st.subheader("債務者データ取得")
            
            # データ取得方法の選択
            data_source = st.radio(
                "データ取得方法を選択してください",
                ["スプレッドシート一覧から選択", "スプレッドシートリンクを直接入力"],
                horizontal=True
            )
            
            selected_debtor = None
            creditor_data = None
            
            if data_source == "スプレッドシート一覧から選択":
                with st.spinner("債務者一覧を取得中..."):
                    spreadsheets = sheets_manager.list_spreadsheets()
                
                if not spreadsheets:
                    st.warning("債務者のスプレッドシートが見つかりません")
                else:
                    debtor_names = [sheet['name'] for sheet in spreadsheets]
                    selected_debtor = st.selectbox("債務者を選択", debtor_names)
                    
                    if selected_debtor:
                        with st.spinner(f"{selected_debtor}のデータを取得中..."):
                            selected_sheet = next(sheet for sheet in spreadsheets if sheet['name'] == selected_debtor)
                            data = safe_get_data_from_sheet_info(sheets_manager, selected_sheet)
                        
                        if data is not None:
                            if isinstance(data, pd.DataFrame) and not data.empty:
                                headers, creditor_data = handle_dataframe_conversion(data)
                                if creditor_data:
                                    st.success("データを取得しました")
                                    with st.expander("データプレビュー"):
                                        df = pd.DataFrame(creditor_data)
                                        st.dataframe(df, use_container_width=True)
                            elif isinstance(data, list) and len(data) > 1:
                                headers, creditor_data = handle_dataframe_conversion(data)
                                if creditor_data:
                                    st.success("データを取得しました")
                                    with st.expander("データプレビュー"):
                                        df = pd.DataFrame(creditor_data)
                                        st.dataframe(df, use_container_width=True)
                            else:
                                st.warning("データが見つかりませんでした")
                        else:
                            st.warning("データが見つかりませんでした")
            
            else:  # スプレッドシートリンクを直接入力
                st.write("**スプレッドシートリンクから直接データを取得**")
                
                # スプレッドシートリンク入力
                spreadsheet_url = st.text_input(
                    "スプレッドシートのリンクを貼り付けてください",
                    placeholder="https://docs.google.com/spreadsheets/d/...",
                    help="Google SheetsのURLを貼り付けてください"
                )
                
                # 債務者名入力（オプション）
                manual_debtor_name = st.text_input(
                    "債務者名（オプション - 空白の場合はスプレッドシートから自動取得）", 
                    placeholder="例：株式会社サンプル",
                    help="入力しない場合は、スプレッドシートのタイトルから自動で抽出します"
                )
                
                if spreadsheet_url:
                    if st.button("データを取得", type="secondary"):
                        with st.spinner("スプレッドシートからデータを取得中..."):
                            try:
                                # URLからスプレッドシートIDを抽出
                                match = re.search(r'/spreadsheets/d/([a-zA-Z0-9-_]+)', spreadsheet_url)
                                if not match:
                                    st.error("有効なGoogle SheetsのURLではありません")
                                else:
                                    spreadsheet_id = match.group(1)
                                    
                                    # スプレッドシート情報を取得
                                    try:
                                        spreadsheet = sheets_manager.gc.open_by_key(spreadsheet_id)
                                        
                                        # 債務者名を決定
                                        if manual_debtor_name.strip():
                                            selected_debtor = manual_debtor_name.strip()
                                        else:
                                            sheet_title = spreadsheet.title
                                            if "債権者データ_" in sheet_title:
                                                parts = sheet_title.split('_')
                                                if len(parts) >= 2:
                                                    selected_debtor = parts[1]
                                                else:
                                                    selected_debtor = "不明な債務者"
                                            else:
                                                selected_debtor = sheet_title
                                        
                                        # セッション状態に保存
                                        st.session_state.selected_debtor = selected_debtor
                                        
                                        # データ取得
                                        data = safe_get_spreadsheet_data_by_id(sheets_manager, spreadsheet_id)
                                        
                                        if data is not None and len(data) > 1:
                                            headers, creditor_data = handle_dataframe_conversion(data)
                                            if creditor_data:
                                                st.session_state.creditor_data = creditor_data
                                                st.success("データを取得しました")
                                                
                                                with st.expander("データプレビュー"):
                                                    df = pd.DataFrame(creditor_data)
                                                    st.dataframe(df, use_container_width=True)
                                            else:
                                                st.error("データの変換に失敗しました")
                                        else:
                                            st.error("スプレッドシートにデータが見つかりません")
                                    
                                    except Exception as sheet_error:
                                        st.error(f"スプレッドシートアクセスエラー: {sheet_error}")
                                        st.write("スプレッドシートが共有されているか、URLが正しいか確認してください")
                            
                            except Exception as e:
                                st.error(f"データ取得エラー: {e}")
                
                # セッション状態から値を復元
                if 'selected_debtor' in st.session_state:
                    selected_debtor = st.session_state.selected_debtor
                if 'creditor_data' in st.session_state:
                    creditor_data = st.session_state.creditor_data
            
            # データが取得できた場合のダウンロード機能
            if selected_debtor and creditor_data:
                st.markdown("---")
                st.subheader("エクスポート")
                
                output_filename = st.text_input("出力ファイル名", 
                    value=f"{datetime.now().strftime('%Y%m%d')}_{selected_debtor}_{procedure_type}_債権者一覧表")
                
                if st.button("債権者一覧表をダウンロード", type="primary", use_container_width=True):
                    with st.spinner("債権者一覧表を作成中..."):
                        try:
                            template_path = template_manager.get_template_path(template_key)
                            file_extension = get_file_extension(template_path)
                            
                            # ファイル形式に応じた処理
                            if file_extension == ".docx":
                                # Word文書の処理
                                processed_doc = process_word_template(
                                    template_path, creditor_data, selected_debtor, selected_court, procedure_type, case_number
                                )
                                
                                output = io.BytesIO()
                                processed_doc.save(output)
                                output.seek(0)
                                
                                mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                file_ext = "docx"
                                
                            else:
                                # Excel文書の処理
                                processed_wb = process_excel_template(
                                    template_path, creditor_data, selected_debtor, selected_court, procedure_type, case_number
                                )
                                
                                output = io.BytesIO()
                                processed_wb.save(output)
                                output.seek(0)
                                
                                mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                                file_ext = "xlsx"
                            
                            # 作成完了メッセージ
                            format_name = "Word" if file_extension == ".docx" else "Excel"
                            st.markdown(get_success_html(f"{procedure_type}の債権者一覧表が作成されました ({format_name}形式)"), unsafe_allow_html=True)
                            
                            # ダウンロードボタン
                            st.download_button(
                                label=f"ダウンロード ({format_name}形式)",
                                data=output.getvalue(),
                                file_name=f"{output_filename}.{file_ext}",
                                mime=mime_type,
                                use_container_width=True,
                                type="secondary"
                            )
                            
                            # ファイル情報
                            file_size = len(output.getvalue())
                            st.caption(f"ファイルサイズ: {file_size:,} バイト ({file_size/1024/1024:.2f} MB)")
                            
                        except Exception as e:
                            st.error(f"処理エラー: {e}")
                            st.write("エラー詳細:")
                            import traceback
                            st.text(traceback.format_exc())
        
        else:
            st.warning(f"{selected_court} - {procedure_type} のテンプレートが登録されていません")
    
    # テンプレート登録タブ
    with tab2:
        st.subheader("債権者一覧表テンプレート登録")
        
        col1, col2 = st.columns(2)
        
        with col1:
            # 裁判所選択
            selected_court_reg = st.selectbox("裁判所を選択", courts, key="court_registration")
            if selected_court_reg == "その他":
                selected_court_reg = st.text_input("裁判所名を入力", key="court_name_registration")
        
        with col2:
            # 手続種別選択
            procedure_type_reg = st.selectbox("手続種別を選択", procedure_types, key="procedure_registration")
        
        template_key_reg = get_template_key(selected_court_reg, procedure_type_reg)
        
        st.write(f"**{selected_court_reg} - {procedure_type_reg}** の債権者一覧表テンプレートを登録")
        
        new_desc = st.text_input("説明", value=f"{procedure_type_reg}用債権者一覧表", placeholder="テンプレートの用途や特徴")
        new_file = st.file_uploader("テンプレートファイル", type=['xlsx', 'docx'], help="ExcelファイルまたはWordファイルを選択してください")
        
        if st.button("テンプレート登録") and new_file:
            # ファイル拡張子をチェック
            file_extension = os.path.splitext(new_file.name)[1].lower()
            if file_extension in ['.xlsx', '.docx']:
                if template_manager.save_template(template_key_reg, new_file.read(), new_desc, file_extension):
                    format_name = "Excel" if file_extension == ".xlsx" else "Word"
                    st.success(f"{selected_court_reg} - {procedure_type_reg} の債権者一覧表テンプレートを登録しました ({format_name}形式)")
                    st.rerun()
            else:
                st.error("Excelファイル(.xlsx)またはWordファイル(.docx)のみ対応しています")
        
        # 既存テンプレート一覧表示
        st.markdown("---")
        st.subheader("登録済みテンプレート一覧")
        
        registered_templates = []
        for court in courts[:-1]:  # "その他"を除く
            for proc_type in procedure_types:
                key = get_template_key(court, proc_type)
                if template_manager.template_exists(key):
                    template_info = template_manager.get_template_info(key)
                    template_path = template_manager.get_template_path(key)
                    file_ext = get_file_extension(template_path)
                    format_name = "Excel" if file_ext == ".xlsx" else "Word" if file_ext == ".docx" else "不明"
                    
                    registered_templates.append({
                        "裁判所": court,
                        "手続種別": proc_type,
                        "形式": format_name,
                        "説明": template_info['description'],
                        "最終更新": template_info['last_modified']
                    })
        
        if registered_templates:
            df_templates = pd.DataFrame(registered_templates)
            st.dataframe(df_templates, use_container_width=True)
        else:
            st.info("登録済みテンプレートはありません")
        
        # テンプレート更新機能
        if template_manager.template_exists(template_key_reg):
            st.markdown("---")
            st.write(f"**{selected_court_reg} - {procedure_type_reg}** の既存テンプレートを更新")
            if st.checkbox("テンプレートを更新"):
                updated_file = st.file_uploader("新しいファイル", type=['xlsx', 'docx'], key="update")
                updated_desc = st.text_input("更新説明", value=new_desc, key="update_desc")
                if updated_file and st.button("更新実行"):
                    file_extension = os.path.splitext(updated_file.name)[1].lower()
                    if file_extension in ['.xlsx', '.docx']:
                        if template_manager.save_template(template_key_reg, updated_file.read(), updated_desc, file_extension):
                            st.success("テンプレートを更新しました")
                            st.rerun()
                    else:
                        st.error("Excelファイル(.xlsx)またはWordファイル(.docx)のみ対応しています")
        
        # テンプレート変数説明
        st.markdown("---")
        with st.expander("テンプレート変数一覧"):
            for category, variables in TEMPLATE_VARIABLES.items():
                st.write(f"**{category}**")
                for var, desc in variables.items():
                    st.write(f"`{var}` → {desc}")
                st.write("")
            
            st.write("**記載例:**")
            st.code("""
債務者: {debtor_name}
裁判所: {court_name}
手続種別: {procedure_type}
事件番号: {case_number}

債権者1: {company_name_1}
支店名: {branch_name_1}
住所: {postal_code_1} {address_1}
電話: {phone_number_1}
債権名: {claim_name_1}
債権額: {claim_amount_1}円
契約日: {contract_date_1}
原債権者: {original_creditor_1}

債権者2: {company_name_2}
住所: {postal_code_2} {address_2}
債権額: {claim_amount_2}円

合計金額: {sum_claim_amount_1_to_2}円
            """)
    
    # 使用方法
    st.markdown("---")
    st.subheader("使用方法")
    st.markdown("""
    **1. テンプレート準備**
    - 各裁判所の個人再生・自己破産用債権者一覧表を用意（Excel または Word）
    - セルまたは文書内に変数を記載（例: {company_name_1}、{procedure_type}）
    
    **2. テンプレート登録**
    - 「テンプレート登録」タブで裁判所・手続種別を選択
    - 対応するExcelファイル(.xlsx)またはWordファイル(.docx)をアップロード
    
    **3. 債権者一覧表作成**
    - 「テンプレート使用」タブで裁判所・手続種別・債務者を選択
    - 「債権者一覧表をダウンロード」で完成版をダウンロード
    
    **対応形式**
    - Excel形式（.xlsx）: セル内の変数を置換
    - Word形式（.docx）:- Word形式（.docx）: 文書内の変数を置換（段落、表、ヘッダー、フッター対応）
   
   **ポイント**
   - 各裁判所ごとに個人再生・自己破産の2つのテンプレートを登録可能
   - ExcelとWordの混在も可能（裁判所によって形式を使い分け）
   - 手続種別に応じて適切な書式が自動選択されます
   """)
       
except Exception as e:
   st.error(f"エクスポート機能の初期化エラー: {e}")
   st.write("詳細なエラー情報:")
   import traceback
   st.text(traceback.format_exc())
   
   st.markdown("---")
   st.subheader("トラブルシューティング")
   st.write("1. Google Sheetsの認証情報が正しく設定されているか確認")
   st.write("2. utils/ディレクトリのファイルが正常に読み込めるか確認")
   st.write("3. 必要なライブラリがインストールされているか確認")
