import io
import os
from datetime import datetime
from openpyxl import load_workbook
from docx import Document

class TemplateProcessor:
    def __init__(self, template_manager):
        self.template_manager = template_manager
    
    def is_tokyo_district_jiko_hasan(self, court_name, procedure_type):
        """東京地裁の自己破産かどうか判定"""
        return court_name == "東京地方裁判所" and procedure_type == "自己破産"
    
    def replace_template_variables(self, text, creditor_data, debtor_name, court_name, procedure_type, case_number=""):
        """テンプレート変数を実際のデータで置換"""
        if not isinstance(text, str):
            return text
            
        result = text
        
        # 基本情報の置換
        result = result.replace("{debtor_name}", str(debtor_name))
        result = result.replace("{court_name}", str(court_name))
        result = result.replace("{case_number}", str(case_number))
        result = result.replace("{procedure_type}", str(procedure_type))
        result = result.replace("{today}", datetime.now().strftime('%Y年%m月%d日'))
        result = result.replace("{today_slash}", datetime.now().strftime('%Y/%m/%d'))
        result = result.replace("{total_creditors}", str(len(creditor_data)))
        
        # 債権金額総計の計算
        total_amount = sum(float(str(row.get('債権額', 0)).replace(',', '')) if row.get('債権額') else 0 for row in creditor_data)
        result = result.replace("{total_claim_amount}", f"{int(total_amount):,}")
        
        # 東京地裁自己破産の特殊処理
        if self.is_tokyo_district_jiko_hasan(court_name, procedure_type):
            result = self._replace_tokyo_variables(result, creditor_data)
        else:
            # 従来の処理（通常の変数置換）
            result = self._replace_standard_variables(result, creditor_data)
        
        return result
    
    def _convert_claim_name_to_code(self, claim_name):
        """東京地裁用：債権名をコードに変換"""
        claim_name_str = str(claim_name).strip()
        conversion_map = {
            '貸付金': 'A',
            '立替金': 'B', 
            '保証金': 'C',
            'その他': 'D'
        }
        return conversion_map.get(claim_name_str, 'D')  # 該当なしの場合はDを返す
    
    def _replace_tokyo_variables(self, text, creditor_data):
        """東京地裁用のA/B変数置換"""
        result = text
        creditor_count = len(creditor_data)
        
        # 債権者数に応じて使用する変数プレフィックスを決定
        # 8人以下：A系変数、9人以上：B系変数
        if creditor_count <= 8:
            prefix = "A"
        else:
            prefix = "B"
        
        # 債権者個別情報の置換（スプレッドシート順を維持）
        for i, creditor in enumerate(creditor_data, 1):
            # 東京地裁専用：債権名をコードに変換
            claim_name_code = self._convert_claim_name_to_code(creditor.get('債権名', ''))
            
            # A系またはB系の変数を置換
            replacements = {
                f"{{company_name_{prefix}{i}}}": str(creditor.get('会社名', '')),
                f"{{branch_name_{prefix}{i}}}": str(creditor.get('支店名', '')),
                f"{{postal_code_{prefix}{i}}}": str(creditor.get('郵便番号', '')),
                f"{{address_{prefix}{i}}}": str(creditor.get('住所', '')),
                f"{{phone_number_{prefix}{i}}}": str(creditor.get('電話番号', '')),
                f"{{fax_number_{prefix}{i}}}": str(creditor.get('FAX番号', '')),
                f"{{claim_name_{prefix}{i}}}": claim_name_code,  # コード変換済み
                f"{{claim_amount_{prefix}{i}}}": str(creditor.get('債権額', '')),
                f"{{contract_date_{prefix}{i}}}": str(creditor.get('契約日', '')),
                f"{{first_borrowing_date_{prefix}{i}}}": str(creditor.get('初回借入日', '')),
                f"{{last_borrowing_date_{prefix}{i}}}": str(creditor.get('最終借入日', '')),
                f"{{last_payment_date_{prefix}{i}}}": str(creditor.get('最終返済日', '')),
                f"{{original_creditor_{prefix}{i}}}": str(creditor.get('原債権者', '')),
                f"{{substitution_or_transfer_{prefix}{i}}}": str(creditor.get('代位弁済/債権譲渡', '')),
                f"{{transfer_date_{prefix}{i}}}": str(creditor.get('債権移転日', '')),
                f"{{status_{prefix}{i}}}": str(creditor.get('ステータス', '')),
                f"{{notes_{prefix}{i}}}": str(creditor.get('備考', '')),
                f"{{registration_date_{prefix}{i}}}": str(creditor.get('登録日', '')),
                f"{{creditor_rank_{prefix}{i}}}": str(i),
                f"{{id_{prefix}{i}}}": str(creditor.get('ID', ''))
            }
            
            for var, value in replacements.items():
                result = result.replace(var, value)
        
        # 使用されなかった方の変数（A系またはB系）を空文字で置換
        unused_prefix = "B" if prefix == "A" else "A"
        for i in range(1, 21):  # 最大20人まで対応
            unused_replacements = {
                f"{{company_name_{unused_prefix}{i}}}": "",
                f"{{branch_name_{unused_prefix}{i}}}": "",
                f"{{postal_code_{unused_prefix}{i}}}": "",
                f"{{address_{unused_prefix}{i}}}": "",
                f"{{phone_number_{unused_prefix}{i}}}": "",
                f"{{fax_number_{unused_prefix}{i}}}": "",
                f"{{claim_name_{unused_prefix}{i}}}": "",
                f"{{claim_amount_{unused_prefix}{i}}}": "",
                f"{{contract_date_{unused_prefix}{i}}}": "",
                f"{{first_borrowing_date_{unused_prefix}{i}}}": "",
                f"{{last_borrowing_date_{unused_prefix}{i}}}": "",
                f"{{last_payment_date_{unused_prefix}{i}}}": "",
                f"{{original_creditor_{unused_prefix}{i}}}": "",
                f"{{substitution_or_transfer_{unused_prefix}{i}}}": "",
                f"{{transfer_date_{unused_prefix}{i}}}": "",
                f"{{status_{unused_prefix}{i}}}": "",
                f"{{notes_{unused_prefix}{i}}}": "",
                f"{{registration_date_{unused_prefix}{i}}}": "",
                f"{{creditor_rank_{unused_prefix}{i}}}": "",
                f"{{id_{unused_prefix}{i}}}": ""
            }
            
            for var, value in unused_replacements.items():
                result = result.replace(var, value)
        
        return result
    
    def _replace_standard_variables(self, text, creditor_data):
        """従来の標準変数置換（他の裁判所用）"""
        result = text
        
        # 債権者個別情報の置換
        for i, creditor in enumerate(creditor_data, 1):
            replacements = {
                f"{{id_{i}}}": str(creditor.get('ID', '')),
                f"{{company_name_{i}}}": str(creditor.get('会社名', '')),
                f"{{branch_name_{i}}}": str(creditor.get('支店名', '')),
                f"{{postal_code_{i}}}": str(creditor.get('郵便番号', '')),
                f"{{address_{i}}}": str(creditor.get('住所', '')),
                f"{{phone_number_{i}}}": str(creditor.get('電話番号', '')),
                f"{{fax_number_{i}}}": str(creditor.get('FAX番号', '')),
                f"{{claim_name_{i}}}": str(creditor.get('債権名', '')),
                f"{{claim_amount_{i}}}": str(creditor.get('債権額', '')),
                f"{{contract_date_{i}}}": str(creditor.get('契約日', '')),
                f"{{first_borrowing_date_{i}}}": str(creditor.get('初回借入日', '')),
                f"{{last_borrowing_date_{i}}}": str(creditor.get('最終借入日', '')),
                f"{{last_payment_date_{i}}}": str(creditor.get('最終返済日', '')),
                f"{{original_creditor_{i}}}": str(creditor.get('原債権者', '')),
                f"{{substitution_or_transfer_{i}}}": str(creditor.get('代位弁済/債権譲渡', '')),
                f"{{transfer_date_{i}}}": str(creditor.get('債権移転日', '')),
                f"{{status_{i}}}": str(creditor.get('ステータス', '')),
                f"{{notes_{i}}}": str(creditor.get('備考', '')),
                f"{{registration_date_{i}}}": str(creditor.get('登録日', '')),
                f"{{creditor_rank_{i}}}": str(i)
            }
            
            for var, value in replacements.items():
                result = result.replace(var, value)
        
        return result
        
    def process_excel_template(self, template_path, creditor_data, debtor_name, court_name, procedure_type, case_number=""):
        """Excelテンプレートファイルを処理"""
        wb = load_workbook(template_path)
        
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value:
                        cell.value = self.replace_template_variables(
                            cell.value, creditor_data, debtor_name, court_name, procedure_type, case_number
                        )
        
        return wb
    
    def process_word_template(self, template_path, creditor_data, debtor_name, court_name, procedure_type, case_number=""):
        """Wordテンプレートファイルを処理"""
        doc = Document(template_path)
        
        # 段落内のテキストを処理
        for paragraph in doc.paragraphs:
            if paragraph.text:
                new_text = self.replace_template_variables(
                    paragraph.text, creditor_data, debtor_name, court_name, procedure_type, case_number
                )
                if new_text != paragraph.text:
                    paragraph.clear()
                    paragraph.add_run(new_text)
        
        # テーブル内のテキストを処理
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text:
                            new_text = self.replace_template_variables(
                                paragraph.text, creditor_data, debtor_name, court_name, procedure_type, case_number
                            )
                            if new_text != paragraph.text:
                                paragraph.clear()
                                paragraph.add_run(new_text)
        
        return doc
    
    def process_template(self, template_key, creditor_data, debtor_name, court_name, procedure_type, case_number=""):
        """テンプレートを処理してファイルを生成"""
        template_path = self.template_manager.get_template_path(template_key)
        file_extension = self.get_file_extension(template_path)
        
        # ファイル形式に応じた処理
        if file_extension == ".docx":
            # Word文書の処理
            processed_doc = self.process_word_template(
                template_path, creditor_data, debtor_name, court_name, procedure_type, case_number
            )
            
            output = io.BytesIO()
            processed_doc.save(output)
            output.seek(0)
            
            return output, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx", "Word"
            
        else:
            # Excel文書の処理
            processed_wb = self.process_excel_template(
                template_path, creditor_data, debtor_name, court_name, procedure_type, case_number
            )
            
            output = io.BytesIO()
            processed_wb.save(output)
            output.seek(0)
            
            return output, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "xlsx", "Excel"
    
    @staticmethod
    def get_file_extension(file_path):
        """ファイル拡張子を取得"""
        return os.path.splitext(file_path)[1].lower()
    
    @staticmethod
    def get_template_key(court_name, procedure_type):
        """テンプレートキーを生成"""
        return f"{court_name}_{procedure_type}"